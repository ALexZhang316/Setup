#!/usr/bin/env python3
"""Audit DOCX structure and enforce common text-only delivery requirements."""

from __future__ import annotations

import argparse
import json
import re
import sys
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any
from zipfile import BadZipFile, ZipFile

from docx import Document


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
}
W = f"{{{NS['w']}}}"
M = f"{{{NS['m']}}}"


def parse_expected(spec: str) -> tuple[str, int]:
    try:
        expression, count = spec.rsplit("=", 1)
        return expression, int(count)
    except (ValueError, TypeError) as exc:
        raise argparse.ArgumentTypeError("expectation must be TEXT_OR_REGEX=COUNT") from exc


def visible_xml_text(data: bytes) -> str:
    root = ET.fromstring(data)
    return "".join(node.text or "" for node in root.iter(f"{W}t")).strip()


def xml_metrics(data: bytes) -> dict[str, Any]:
    root = ET.fromstring(data)
    fields = [node.text or "" for node in root.iter(f"{W}instrText")]
    fields.extend(node.attrib.get(f"{W}instr", "") for node in root.iter(f"{W}fldSimple"))
    return {
        "drawing": sum(1 for _ in root.iter(f"{W}drawing")),
        "pict": sum(1 for _ in root.iter(f"{W}pict")),
        "object": sum(1 for _ in root.iter(f"{W}object")),
        "math": sum(1 for _ in root.iter(f"{M}oMath")),
        "math_paragraph": sum(1 for _ in root.iter(f"{M}oMathPara")),
        "fields": [field.strip() for field in fields if field.strip()],
    }


def run_summary(run) -> dict[str, Any]:
    return {
        "text": run.text,
        "bold": run.bold,
        "italic": run.italic,
        "font": run.font.name,
        "size_pt": run.font.size.pt if run.font.size else None,
    }


def paragraph_summary(paragraph) -> dict[str, Any]:
    paragraph_format = paragraph.paragraph_format
    return {
        "text": paragraph.text,
        "style": paragraph.style.name if paragraph.style else None,
        "space_before_pt": paragraph_format.space_before.pt if paragraph_format.space_before else None,
        "space_after_pt": paragraph_format.space_after.pt if paragraph_format.space_after else None,
        "line_spacing": str(paragraph_format.line_spacing) if paragraph_format.line_spacing else None,
        "runs": [run_summary(run) for run in paragraph.runs],
    }


def length_value(value) -> int | None:
    return int(value) if value is not None else None


def audit(path: Path, sample_count: int) -> tuple[dict[str, Any], str]:
    document = Document(path)
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    nonempty = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]

    with ZipFile(path) as archive:
        bad_member = archive.testzip()
        names = archive.namelist()
        media = sorted(name for name in names if name.startswith("word/media/"))
        embeddings = sorted(name for name in names if name.startswith("word/embeddings/"))
        header_footer_text: dict[str, str] = {}
        totals = {"drawing": 0, "pict": 0, "object": 0, "math": 0, "math_paragraph": 0}
        fields: list[dict[str, str]] = []
        xml_errors: list[str] = []

        for name in names:
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            try:
                data = archive.read(name)
                metrics = xml_metrics(data)
                for key in totals:
                    totals[key] += metrics[key]
                fields.extend({"part": name, "instruction": value} for value in metrics["fields"])
                if re.fullmatch(r"word/(?:header|footer)\d+\.xml", name):
                    header_footer_text[name] = visible_xml_text(data)
            except ET.ParseError as exc:
                xml_errors.append(f"{name}: {exc}")

    report: dict[str, Any] = {
        "schema_version": 1,
        "path": str(path.resolve()),
        "size_bytes": path.stat().st_size,
        "zip_integrity": bad_member is None,
        "bad_zip_member": bad_member,
        "paragraph_count": len(document.paragraphs),
        "nonempty_paragraph_count": len(nonempty),
        "table_count": len(document.tables),
        "section_count": len(document.sections),
        "sections": [
            {
                "orientation": str(section.orientation),
                "page_width": length_value(section.page_width),
                "page_height": length_value(section.page_height),
                "top_margin": length_value(section.top_margin),
                "bottom_margin": length_value(section.bottom_margin),
                "left_margin": length_value(section.left_margin),
                "right_margin": length_value(section.right_margin),
            }
            for section in document.sections
        ],
        "media": media,
        "embeddings": embeddings,
        "xml_objects": totals,
        "field_instructions": fields,
        "header_footer_text": header_footer_text,
        "xml_errors": xml_errors,
        "sample_paragraphs": [paragraph_summary(paragraph) for paragraph in nonempty[:sample_count]],
    }
    return report, full_text


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Inspect and validate DOCX package structure.")
    parser.add_argument("docx", type=Path, help="DOCX file to audit")
    parser.add_argument("--json", dest="json_path", type=Path, help="Write JSON report to this path")
    parser.add_argument("--sample-paragraphs", type=int, default=5, help="Include this many nonempty paragraphs")
    parser.add_argument("--count-regex", action="append", default=[], help="Report a regex match count")
    parser.add_argument(
        "--expect-regex",
        action="append",
        default=[],
        type=parse_expected,
        metavar="REGEX=COUNT",
        help="Require a regex match count",
    )
    parser.add_argument(
        "--expect-marker",
        action="append",
        default=[],
        type=parse_expected,
        metavar="TEXT=COUNT",
        help="Require an exact marker occurrence count",
    )
    parser.add_argument("--expect-text-only", action="store_true", help="Reject media, drawings, objects, and math")
    parser.add_argument("--expect-no-page-fields", action="store_true", help="Reject PAGE/NUMPAGES fields")
    parser.add_argument("--expect-no-header-footer-text", action="store_true", help="Reject visible header/footer text")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    source = args.docx.resolve()
    if not source.exists():
        print(f"error: file not found: {source}", file=sys.stderr)
        return 2

    try:
        report, full_text = audit(source, args.sample_paragraphs)
    except (BadZipFile, KeyError, ValueError) as exc:
        print(f"error: invalid DOCX: {type(exc).__name__}: {exc}", file=sys.stderr)
        return 2

    failures: list[str] = []
    regex_counts: dict[str, int] = {}
    for expression in args.count_regex:
        try:
            regex_counts[expression] = len(re.findall(expression, full_text, flags=re.MULTILINE))
        except re.error as exc:
            failures.append(f"invalid count regex {expression!r}: {exc}")

    for expression, expected in args.expect_regex:
        try:
            actual = len(re.findall(expression, full_text, flags=re.MULTILINE))
            regex_counts[expression] = actual
            if actual != expected:
                failures.append(f"regex {expression!r}: expected {expected}, found {actual}")
        except re.error as exc:
            failures.append(f"invalid expectation regex {expression!r}: {exc}")

    marker_counts: dict[str, int] = {}
    for marker, expected in args.expect_marker:
        actual = full_text.count(marker)
        marker_counts[marker] = actual
        if actual != expected:
            failures.append(f"marker {marker!r}: expected {expected}, found {actual}")

    if args.expect_text_only:
        if report["media"]:
            failures.append(f"media parts present: {len(report['media'])}")
        if report["embeddings"]:
            failures.append(f"embedded objects present: {len(report['embeddings'])}")
        for key, count in report["xml_objects"].items():
            if count:
                failures.append(f"OOXML {key} elements present: {count}")

    if args.expect_no_page_fields:
        page_fields = [
            field
            for field in report["field_instructions"]
            if re.search(r"\b(?:PAGE|NUMPAGES|SECTIONPAGES)\b", field["instruction"], re.IGNORECASE)
        ]
        if page_fields:
            failures.append(f"page fields present: {len(page_fields)}")

    if args.expect_no_header_footer_text:
        populated = {name: text for name, text in report["header_footer_text"].items() if text.strip()}
        if populated:
            failures.append(f"visible header/footer text present in: {sorted(populated)}")

    if not report["zip_integrity"]:
        failures.append(f"ZIP integrity failed at {report['bad_zip_member']}")
    if report["xml_errors"]:
        failures.append(f"XML parse errors: {len(report['xml_errors'])}")

    report["regex_counts"] = regex_counts
    report["marker_counts"] = marker_counts
    report["expectations"] = {"passed": not failures, "failures": failures}
    output = json.dumps(report, ensure_ascii=False, indent=2) + "\n"
    if args.json_path:
        args.json_path.parent.mkdir(parents=True, exist_ok=True)
        args.json_path.write_text(output, encoding="utf-8")
        print(str(args.json_path.resolve()))
    else:
        print(output, end="")
    return 0 if not failures else 2


if __name__ == "__main__":
    sys.exit(main())

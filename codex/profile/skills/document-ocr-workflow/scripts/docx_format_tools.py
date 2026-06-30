#!/usr/bin/env python3
"""OOXML-safe helpers for transplanting DOCX paragraph and run formatting."""

from __future__ import annotations

import argparse
import copy
import os
import re
import sys
import tempfile
import uuid
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def clear_paragraph_content(paragraph: Paragraph) -> None:
    """Remove paragraph content while retaining paragraph properties."""
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def clear_body_keep_section(document: Document) -> None:
    """Remove body content while keeping the final section-properties element."""
    body = document._element.body
    section_properties = body.sectPr
    for child in list(body):
        if child is not section_properties:
            body.remove(child)


def clone_paragraph(document: Document, template: Paragraph, clear_content: bool = True) -> Paragraph:
    """Insert a cloned paragraph before body sectPr and optionally clear its text."""
    body = document._element.body
    paragraph_xml = copy.deepcopy(template._p)
    section_properties = body.sectPr
    if section_properties is None:
        body.append(paragraph_xml)
    else:
        body.insert(body.index(section_properties), paragraph_xml)
    paragraph = Paragraph(paragraph_xml, body)
    if clear_content:
        clear_paragraph_content(paragraph)
    return paragraph


def copy_run_format(target: Run, source: Run) -> None:
    """Replace target run properties with a deep copy of source run properties."""
    if target._r.rPr is not None:
        target._r.remove(target._r.rPr)
    if source._r.rPr is not None:
        target._r.insert(0, copy.deepcopy(source._r.rPr))


def add_formatted_run(paragraph: Paragraph, text: str, source: Run) -> Run:
    """Add text using the exact run properties of a donor run."""
    run = paragraph.add_run(text)
    copy_run_format(run, source)
    return run


def add_runs_from_donors(
    paragraph: Paragraph,
    segments: Iterable[tuple[str, Run]],
) -> list[Run]:
    """Add multiple text segments using separate donor runs."""
    return [add_formatted_run(paragraph, text, donor) for text, donor in segments]


def find_paragraph(document: Document, pattern: str, flags: int = 0) -> Paragraph:
    """Find the first paragraph whose visible text matches a regular expression."""
    expression = re.compile(pattern, flags)
    for paragraph in document.paragraphs:
        if expression.search(paragraph.text):
            return paragraph
    raise LookupError(f"no paragraph matches: {pattern}")


def copy_section_geometry(source, target) -> None:
    """Copy page geometry and section distances without copying headers or footers."""
    attributes = (
        "orientation",
        "page_width",
        "page_height",
        "top_margin",
        "bottom_margin",
        "left_margin",
        "right_margin",
        "header_distance",
        "footer_distance",
        "gutter",
        "start_type",
    )
    for attribute in attributes:
        value = getattr(source, attribute, None)
        if value is not None:
            setattr(target, attribute, value)


def clear_headers_footers(document: Document) -> None:
    """Clear all existing header/footer parts without creating new story parts."""
    for relationship in document.part.rels.values():
        if relationship.reltype not in (RT.HEADER, RT.FOOTER):
            continue
        root = relationship.target_part.element
        for child in list(root):
            root.remove(child)
        root.append(OxmlElement("w:p"))


def safe_save(document: Document, destination: str | Path) -> Path:
    """Save beside the destination and atomically replace it, preserving a lock-safe temp file."""
    destination = Path(destination).resolve()
    destination.parent.mkdir(parents=True, exist_ok=True)
    temporary = destination.with_name(
        f".{destination.stem}.{uuid.uuid4().hex}.tmp{destination.suffix}"
    )
    document.save(temporary)
    try:
        os.replace(temporary, destination)
    except PermissionError as exc:
        raise RuntimeError(
            f"destination is probably open in Word: {destination}; "
            f"close it and replace it with the completed temporary file: {temporary}"
        ) from exc
    except Exception:
        temporary.unlink(missing_ok=True)
        raise
    return destination


def self_test() -> None:
    with tempfile.TemporaryDirectory(prefix="docx-format-tools-") as tmp:
        tmp_path = Path(tmp)
        donor_path = tmp_path / "donor.docx"
        output_path = tmp_path / "output.docx"

        donor = Document()
        donor.sections[0].left_margin = Inches(0.75)
        donor.sections[0].right_margin = Inches(0.8)
        donor.sections[0].orientation = WD_ORIENT.PORTRAIT
        donor.sections[0].header.paragraphs[0].text = "REMOVE ME"
        paragraph = donor.add_paragraph()
        prefix = paragraph.add_run("PREFIX\t")
        prefix.bold = True
        prefix.font.name = "Times New Roman"
        prefix.font.size = Pt(9)
        stem = paragraph.add_run("Stem")
        stem.font.name = "Times New Roman"
        stem.font.size = Pt(9)
        donor.save(donor_path)

        donor = Document(donor_path)
        template = donor.paragraphs[-1]
        target = Document()
        copy_section_geometry(donor.sections[0], target.sections[0])
        clear_body_keep_section(target)
        cloned = clone_paragraph(target, template)
        add_runs_from_donors(
            cloned,
            [("ANSWER\t", template.runs[0]), ("New stem", template.runs[1])],
        )
        clear_headers_footers(target)
        safe_save(target, output_path)

        reopened = Document(output_path)
        assert reopened._element.body[-1].tag == qn("w:sectPr")
        assert reopened.paragraphs[0].text == "ANSWER\tNew stem"
        assert reopened.paragraphs[0].runs[0].bold is True
        assert reopened.paragraphs[0].runs[0].font.size.pt == 9
        assert reopened.sections[0].left_margin == donor.sections[0].left_margin
        assert reopened.sections[0].header.paragraphs[0].text == ""
        print(f"self-test passed: {output_path}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Import this module for DOCX format cloning, or run its smoke test."
    )
    parser.add_argument("--self-test", action="store_true", help="Run an isolated format-clone smoke test")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if not args.self_test:
        print("Import this module, or pass --self-test.")
        return 0
    self_test()
    return 0


if __name__ == "__main__":
    sys.exit(main())

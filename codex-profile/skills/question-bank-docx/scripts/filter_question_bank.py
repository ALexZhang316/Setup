from __future__ import annotations

import argparse
import re
import sys
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


QUESTION_RE = re.compile(r"^([A-E])\s*(\d+)[.\uFF0E](.*)$")
GROUP_RE = re.compile(r"^\uFF08(\d+)-(\d+)\u9898\u5171\u7528(\u9898\u5E72|\u9009\u9879)\uFF09$")
OPTION_RE = re.compile(r"^[A-E][.\uFF0E]")

KIND_STEM_TEXT = "\u9898\u5E72"
KIND_STEM = "stem"
KIND_OPTIONS = "options"
LATIN_FONT = "Calibri"
EAST_ASIA_FONT = "\u5fae\u8f6f\u96c5\u9ed1"

DRUG_LITERALS = [
    "\u5229\u5c3f\u5242",
    "\u5f3a\u5fc3\u5242",
    "\u652f\u6c14\u7ba1\u8212\u5f20\u836f",
    "\u652f\u6c14\u7ba1\u8212\u5f20\u5242",
    "\u6297\u751f\u7d20",
    "\u6297\u83cc\u7d20",
    "\u6297\u83cc\u836f\u7269",
    "\u7cd6\u76ae\u8d28\u6fc0\u7d20",
    "\u795b\u75f0\u5242",
    "\u5730\u897f\u6cee",
    "\u6c28\u8336\u78b1",
    "\u5730\u585e\u7c73\u677e",
    "\u5934\u5b62\u66f2\u677e",
    "\u500d\u6c2f\u7c73\u677e",
    "\u767d\u4e09\u70ef\u8c03\u8282\u5242",
    "\u767d\u4e09\u70ef\u53d7\u4f53\u8c03\u8282\u5242",
    "H1\u53d7\u4f53\u62ee\u6297\u5242",
    "H1\u53d7\u4f53\u963b\u65ad\u836f",
    "M\u53d7\u4f53\u62ee\u6297\u5242",
    "M\u53d7\u4f53\u963b\u65ad\u836f",
    "N\u53d7\u4f53\u963b\u65ad\u836f",
    "\u03b1\u53d7\u4f53\u963b\u65ad\u836f",
    "\u03b2\u53d7\u4f53\u963b\u65ad\u836f",
    "\u03b22\u53d7\u4f53\u6fc0\u52a8\u5242",
    "\u03b2\u53d7\u4f53\u963b\u6ede\u5242",
    "\u03b2\u53d7\u4f53\u62ee\u6297\u5242",
    "\u8336\u78b1",
    "\u80be\u4e0a\u817a\u7d20",
    "\u53bb\u7532\u80be\u4e0a\u817a\u7d20",
    "\u5f02\u4e19\u80be\u4e0a\u817a\u7d20",
    "\u6cfc\u5c3c\u677e",
    "\u6c22\u5316\u53ef\u7684\u677e",
    "\u7b2c\u4e09\u4ee3\u5934\u5b62\u83cc\u7d20",
    "\u7532\u57fa\u6cfc\u5c3c\u677e\u9f99",
    "\u9ebb\u9ec4\u7d20",
    "\u9499\u901a\u9053\u963b\u6ede\u5242",
    "\u9752\u9709\u7d20",
    "\u9752\u9709\u7d20G",
    "\u7ea2\u9709\u7d20",
    "\u514b\u62c9\u9709\u7d20",
    "\u5e86\u5927\u9709\u7d20",
    "\u6797\u53ef\u9709\u7d20",
    "\u963f\u5947\u9709\u7d20",
    "\u963f\u83ab\u897f\u6797",
    "\u5de6\u6c27\u6c1f\u6c99\u661f",
    "\u73af\u4e19\u6c99\u661f",
    "\u5934\u5b62\u5458\u8f9b",
    "\u5934\u5b62\u83cc\u7d20\u7c7b",
    "\u9752\u9709\u7d20\u7c7b",
    "\u78b3\u9752\u9709\u70ef\u7c7b",
    "\u5934\u5b62\u5511\u5576",
    "\u514b\u6797\u9709\u7d20",
    "\u4e07\u53e4\u9709\u7d20",
    "\u6c2f\u9709\u7d20",
    "\u4e24\u6027\u9709\u7d20B",
    "\u706d\u6ef4\u7075",
    "\u7ef4\u751f\u7d20B6",
    "\u94fe\u9709\u7d20",
    "\u5229\u798f\u5e73",
    "\u62a4\u809d\u836f",
    "\u4f4e\u5206\u5b50\u91cf\u809d\u7d20",
    "\u666e\u901a\u809d\u7d20",
    "\u94fe\u6fc0\u9176",
    "\u6d0b\u5730\u9ec4",
    "\u975e\u6d0b\u5730\u9ec4\u7c7b\u5f3a\u5fc3\u836f",
    "\u897f\u5730\u5170",
    "\u7f8e\u897f\u5f8b",
    "\u4ed6\u6c40\u7c7b\u836f\u7269",
    "\u6297\u51dd\u836f\u7269",
    "\u6297\u8840\u5c0f\u677f\u836f\u7269",
    "\u6297\u8840\u5c0f\u677f\u836f",
    "\u785d\u9178\u916f\u7c7b\u836f\u7269",
    "\u975e\u753e\u4f53\u6297\u708e\u836f",
    "\u9499\u79bb\u5b50\u62ee\u6297\u5242",
    "\u9499\u62ee\u6297\u5242",
    "\u8840\u7ba1\u7d27\u5f20\u7d20\u8f6c\u6362\u9176\u6291\u5236\u5242",
    "\u8840\u7ba1\u7d27\u7d20\u8f6c\u6362\u9176\u6291\u5236\u5242",
    "\u78b1\u6027\u836f\u7269",
    "\u5347\u538b\u836f",
    "\u6297\u8fc7\u654f\u836f\u7269",
    "\u591a\u5df4\u80fa",
]

IMAGING_REGEXES = [
    ("HRCT", r"(?<![A-Za-z])HRCT(?![A-Za-z])"),
    ("CTPA", r"(?<![A-Za-z])CTPA(?![A-Za-z])"),
    ("PET-CT", r"(?<![A-Za-z])PET-CT(?![A-Za-z])"),
    ("MRI", r"(?<![A-Za-z])MRI(?![A-Za-z])"),
    ("MRCP", r"(?<![A-Za-z])MRCP(?![A-Za-z])"),
    ("ECT", r"(?<![A-Za-z])ECT(?![A-Za-z])"),
    ("CT", r"(?<![A-Za-z])CT(?![A-Za-z])"),
    ("X\u7ebf", r"X\u7ebf"),
    ("\u80f8\u7247", r"\u80f8\u7247"),
    ("\u5e73\u7247", r"\u5e73\u7247"),
    ("\u900f\u89c6", r"\u900f\u89c6"),
    ("\u6444\u7247", r"\u6444\u7247"),
    ("\u8d85\u58f0", r"\u8d85\u58f0"),
    ("B\u8d85", r"B\u8d85"),
    ("\u5f69\u8d85", r"\u5f69\u8d85"),
    ("\u9020\u5f71", r"\u9020\u5f71"),
    ("\u6838\u7d20", r"\u6838\u7d20"),
    ("\u94bc\u9776", r"\u94bc\u9776"),
    ("\u94a1\u9910", r"\u94a1\u9910"),
]


@dataclass
class GroupRecord:
    marker_para_index: int
    marker_line: str
    start_no: int
    end_no: int
    kind: str
    shared_lines: list[str] = field(default_factory=list)
    question_indices: list[int] = field(default_factory=list)


@dataclass
class QuestionRecord:
    source_para_index: int
    answer: str
    question_no: int
    head_line: str
    body_lines: list[str]
    local_option_lines: list[str]
    group_id: int | None
    group_kind: str | None
    shared_stem_lines: list[str]
    shared_option_lines: list[str]
    matched_labels: list[str]


def extract_paragraphs(path: Path) -> list[str]:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml")
    root = ET.fromstring(xml)
    paragraphs: list[str] = []
    for para in root.findall(".//w:body/w:p", ns):
        texts = [node.text for node in para.findall(".//w:t", ns) if node.text]
        text = "".join(texts).strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def parse_question_head(text: str) -> tuple[str, int, str] | None:
    match = QUESTION_RE.match(text)
    if not match:
        return None
    return match.group(1), int(match.group(2)), match.group(3)


def parse_group_marker(text: str) -> tuple[int, int, str] | None:
    match = GROUP_RE.match(text)
    if not match:
        return None
    kind_text = match.group(3)
    kind = KIND_STEM if kind_text == KIND_STEM_TEXT else KIND_OPTIONS
    return int(match.group(1)), int(match.group(2)), kind


def is_option_line(text: str) -> bool:
    return bool(OPTION_RE.match(text))


def normalize_head_line(text: str) -> str:
    parsed = parse_question_head(text)
    if not parsed:
        return text
    answer, question_no, rest = parsed
    return f"{answer} {question_no}.{rest}"


def set_style_fonts(style, latin_font: str = LATIN_FONT, east_asia_font: str = EAST_ASIA_FONT) -> None:
    style.font.name = latin_font
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def configure_document_styles(document: Document) -> None:
    set_style_fonts(document.styles["Normal"])


def compile_matchers(args: argparse.Namespace) -> list[tuple[str, re.Pattern[str]]]:
    matchers: list[tuple[str, re.Pattern[str]]] = []

    if args.preset == "drug":
        for literal in DRUG_LITERALS:
            matchers.append((literal, re.compile(re.escape(literal))))
    elif args.preset == "imaging":
        for label, pattern in IMAGING_REGEXES:
            matchers.append((label, re.compile(pattern)))

    for literal in args.literal:
        matchers.append((literal, re.compile(re.escape(literal))))
    for pattern in args.regex:
        matchers.append((pattern, re.compile(pattern)))

    if not matchers:
        raise SystemExit("no matchers configured; use --preset, --literal, or --regex")
    return matchers


def parse_source(paragraphs: list[str]) -> tuple[list[QuestionRecord], dict[int, GroupRecord]]:
    questions: list[QuestionRecord] = []
    groups: dict[int, GroupRecord] = {}
    pending_group: GroupRecord | None = None
    i = 0

    while i < len(paragraphs):
        group_info = parse_group_marker(paragraphs[i])
        if group_info:
            start_no, end_no, kind = group_info
            shared_lines: list[str] = []
            j = i + 1
            if kind == KIND_STEM:
                while j < len(paragraphs) and not parse_question_head(paragraphs[j]):
                    shared_lines.append(paragraphs[j])
                    j += 1
            else:
                while j < len(paragraphs) and is_option_line(paragraphs[j]):
                    shared_lines.append(paragraphs[j])
                    j += 1
            pending_group = GroupRecord(
                marker_para_index=i,
                marker_line=paragraphs[i],
                start_no=start_no,
                end_no=end_no,
                kind=kind,
                shared_lines=shared_lines,
            )
            groups[pending_group.marker_para_index] = pending_group
            i = j
            continue

        question_info = parse_question_head(paragraphs[i])
        if not question_info:
            i += 1
            continue

        answer, question_no, _ = question_info
        body_lines: list[str] = []
        j = i + 1
        while j < len(paragraphs):
            if parse_group_marker(paragraphs[j]) or parse_question_head(paragraphs[j]):
                break
            body_lines.append(paragraphs[j])
            j += 1

        group_id: int | None = None
        group_kind: str | None = None
        shared_stem_lines: list[str] = []
        shared_option_lines: list[str] = []
        if pending_group and pending_group.start_no <= question_no <= pending_group.end_no:
            group_id = pending_group.marker_para_index
            group_kind = pending_group.kind
            if group_kind == KIND_STEM:
                shared_stem_lines = pending_group.shared_lines
            else:
                shared_option_lines = pending_group.shared_lines

        local_option_lines = [line for line in body_lines if is_option_line(line)]
        question = QuestionRecord(
            source_para_index=i,
            answer=answer,
            question_no=question_no,
            head_line=paragraphs[i],
            body_lines=body_lines,
            local_option_lines=local_option_lines,
            group_id=group_id,
            group_kind=group_kind,
            shared_stem_lines=shared_stem_lines,
            shared_option_lines=shared_option_lines,
            matched_labels=[],
        )
        questions.append(question)
        if group_id is not None:
            groups[group_id].question_indices.append(len(questions) - 1)

        if pending_group and question_no == pending_group.end_no:
            pending_group = None

        i = j

    return questions, groups


def build_search_lines(question: QuestionRecord, match_scope: str) -> list[str]:
    option_lines = question.local_option_lines.copy()
    for line in question.shared_option_lines:
        if line not in option_lines:
            option_lines.append(line)

    if match_scope == "options":
        return option_lines
    if match_scope == "stem_and_options":
        lines = [question.head_line]
        lines.extend(question.shared_stem_lines)
        lines.extend(option_lines)
        return lines
    if match_scope == "full_block":
        lines = [question.head_line]
        lines.extend(question.shared_stem_lines)
        lines.extend(question.body_lines)
        for line in question.shared_option_lines:
            if line not in lines:
                lines.append(line)
        return lines
    raise ValueError(f"unsupported match scope: {match_scope}")


def match_questions(
    questions: list[QuestionRecord],
    groups: dict[int, GroupRecord],
    matchers: list[tuple[str, re.Pattern[str]]],
    match_scope: str,
    expand_shared_stem: bool,
    expand_shared_options: bool,
) -> tuple[set[int], set[int]]:
    direct_hits: set[int] = set()
    for idx, question in enumerate(questions):
        search_lines = build_search_lines(question, match_scope)
        question.matched_labels = [
            label for label, pattern in matchers if any(pattern.search(line) for line in search_lines)
        ]
        if question.matched_labels:
            direct_hits.add(idx)

    selected = set(direct_hits)
    for idx in direct_hits:
        group_id = questions[idx].group_id
        if group_id is None:
            continue
        group = groups[group_id]
        if group.kind == KIND_STEM and expand_shared_stem:
            selected.update(group.question_indices)
        if group.kind == KIND_OPTIONS and expand_shared_options:
            selected.update(group.question_indices)
    return direct_hits, selected


def build_output_units(
    questions: list[QuestionRecord],
    groups: dict[int, GroupRecord],
    selected_question_indices: set[int],
) -> list[tuple[str, int, int]]:
    units: list[tuple[str, int, int]] = []
    emitted_groups: set[int] = set()
    for idx, question in enumerate(questions):
        if idx not in selected_question_indices:
            continue
        if question.group_id is None:
            units.append(("question", question.source_para_index, idx))
            continue
        if question.group_id in emitted_groups:
            continue
        emitted_groups.add(question.group_id)
        units.append(("group", groups[question.group_id].marker_para_index, question.group_id))
    units.sort(key=lambda item: item[1])
    return units


def append_question(document: Document, question: QuestionRecord, group: GroupRecord | None = None) -> None:
    document.add_paragraph(normalize_head_line(question.head_line))
    skip_local_options = (
        group is not None
        and group.kind == KIND_OPTIONS
        and question.local_option_lines
        and question.local_option_lines == group.shared_lines
    )
    for line in question.body_lines:
        if skip_local_options and is_option_line(line):
            continue
        document.add_paragraph(line)


def write_output(
    questions: list[QuestionRecord],
    groups: dict[int, GroupRecord],
    selected_question_indices: set[int],
    output_path: Path,
    title: str,
) -> None:
    document = Document()
    configure_document_styles(document)
    document.add_paragraph(title)

    for unit_type, _, payload in build_output_units(questions, groups, selected_question_indices):
        if unit_type == "question":
            append_question(document, questions[payload])
            continue

        group = groups[payload]
        document.add_paragraph(group.marker_line)
        for line in group.shared_lines:
            document.add_paragraph(line)
        for question_index in group.question_indices:
            if question_index not in selected_question_indices:
                continue
            append_question(document, questions[question_index], group)

    document.save(output_path)


def require_equal(actual: int, expected: int | None, label: str) -> None:
    if expected is None:
        return
    if actual != expected:
        raise SystemExit(f"{label} mismatch: expected {expected}, got {actual}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Filter a structured question-bank .docx and write an answer-first output document."
    )
    parser.add_argument("--source", required=True, help="Path to source .docx")
    parser.add_argument("--output", help="Path to output .docx")
    parser.add_argument("--title", help="Output document title; defaults to output stem")
    parser.add_argument("--preset", choices=["drug", "imaging"], help="Built-in matcher preset")
    parser.add_argument("--literal", action="append", default=[], help="Literal text matcher")
    parser.add_argument("--regex", action="append", default=[], help="Regex matcher")
    parser.add_argument(
        "--match-scope",
        choices=["options", "stem_and_options", "full_block"],
        default="options",
        help="Which lines participate in matching",
    )
    parser.add_argument("--summary-only", action="store_true", help="Print counts without writing output")
    parser.add_argument("--allow-overwrite", action="store_true", help="Allow replacing an existing output file")
    parser.add_argument(
        "--no-expand-shared-stem",
        action="store_true",
        help="Do not expand direct hits to the rest of a shared-stem group",
    )
    parser.add_argument(
        "--no-expand-shared-options",
        action="store_true",
        help="Do not expand direct hits to the rest of a shared-options group",
    )
    parser.add_argument("--sample-count", type=int, default=12, help="How many sample heads to print")
    parser.add_argument("--expected-questions", type=int)
    parser.add_argument("--expected-groups", type=int)
    parser.add_argument("--expected-stem-groups", type=int)
    parser.add_argument("--expected-option-groups", type=int)
    parser.add_argument("--expected-direct", type=int)
    parser.add_argument("--expected-selected", type=int)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    source_path = Path(args.source)
    if not source_path.exists():
        raise SystemExit(f"missing source file: {source_path}")

    matchers = compile_matchers(args)
    paragraphs = extract_paragraphs(source_path)
    questions, groups = parse_source(paragraphs)
    direct_hits, selected_question_indices = match_questions(
        questions,
        groups,
        matchers,
        args.match_scope,
        expand_shared_stem=not args.no_expand_shared_stem,
        expand_shared_options=not args.no_expand_shared_options,
    )

    stem_groups = sum(1 for group in groups.values() if group.kind == KIND_STEM)
    option_groups = sum(1 for group in groups.values() if group.kind == KIND_OPTIONS)
    selected_stem_groups = sum(
        1
        for group in groups.values()
        if group.kind == KIND_STEM and any(index in selected_question_indices for index in group.question_indices)
    )
    selected_option_groups = sum(
        1
        for group in groups.values()
        if group.kind == KIND_OPTIONS and any(index in selected_question_indices for index in group.question_indices)
    )

    require_equal(len(questions), args.expected_questions, "question count")
    require_equal(len(groups), args.expected_groups, "group count")
    require_equal(stem_groups, args.expected_stem_groups, "stem-group count")
    require_equal(option_groups, args.expected_option_groups, "option-group count")
    require_equal(len(direct_hits), args.expected_direct, "direct-hit count")
    require_equal(len(selected_question_indices), args.expected_selected, "selected-question count")

    print(f"source={source_path}")
    print(f"questions={len(questions)}")
    print(f"groups={len(groups)}")
    print(f"stem_groups={stem_groups}")
    print(f"option_groups={option_groups}")
    print(f"direct_hits={len(direct_hits)}")
    print(f"selected_questions={len(selected_question_indices)}")
    print(f"selected_stem_groups={selected_stem_groups}")
    print(f"selected_option_groups={selected_option_groups}")
    print("sample_heads=")
    for idx in sorted(selected_question_indices)[: args.sample_count]:
        print(questions[idx].head_line)

    if args.summary_only:
        return

    if not args.output:
        raise SystemExit("--output is required unless --summary-only is set")

    output_path = Path(args.output)
    if output_path.exists() and not args.allow_overwrite:
        raise SystemExit(f"output file already exists, stopping: {output_path}")

    title = args.title or output_path.stem
    write_output(questions, groups, selected_question_indices, output_path, title)
    print(f"output={output_path}")


if __name__ == "__main__":
    main()
